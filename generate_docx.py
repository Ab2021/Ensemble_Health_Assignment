"""Generate Project_documentation.docx from WRITEUP.md -- manager-facing 2-3 page summary."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def set_cell_shading(cell, color_hex):
    from docx.oxml import parse_xml
    shading = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_para(doc, text, bold=False, italic=False, font_size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p

def add_bullet(doc, text, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p

def build_document():
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Claim Denial Risk Prediction System")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Project Documentation -- Manager Summary")
    run2.italic = True
    run2.font.size = Pt(12)
    doc.add_paragraph()

    # Meta line
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runm = meta.add_run("Ensemble Health Partners | AI Team Take-Home Assessment | May 2026")
    runm.font.size = Pt(10)
    runm.italic = True
    doc.add_paragraph()

    # Section 1: Executive Summary
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(doc,
        "We built a claim denial risk prediction system that identifies which claims are most likely "
        "to be denied before they are submitted. The model ranks all 500 current claims by risk and "
        "assigns each to a High, Medium, or Low tier. The review team can focus the limited 25% "
        "review capacity on the highest-risk claims, catching nearly twice as many denials as "
        "random triage.",
        space_after=8)
    add_para(doc,
        "A GenAI explanation engine produces plain-English summaries for the top 10 highest-risk claims, "
        "telling the analyst exactly why a claim is flagged and what to fix before submission.",
        space_after=8)

    # Section 2: Business Impact
    add_heading(doc, "2. Business Impact", level=1)
    add_para(doc,
        "The model captures 45.7% of all denials within the top 25% of claims -- a 1.76x improvement "
        "over random selection (which would catch only 26%). Of the 139 claims flagged for review, "
        "approximately 48 out of every 100 are actual denials that can be corrected before submission.",
        space_after=8)
    add_para(doc,
        "Based on the historical data, denied claims represent roughly $4.58 million in expected payments. "
        "By catching 45.7% of these in the review queue, the system could help protect approximately "
        "$2.09 million in payments that might otherwise be delayed or lost. Each batch prevents roughly "
        "31 additional denials compared to random triage.",
        space_after=8)

    # Risk tier table
    add_para(doc, "Current Claims Tier Distribution:", bold=True, space_after=4)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Tier"
    hdr[1].text = "Count"
    hdr[2].text = "Review Priority"
    hdr[3].text = "Probability Range"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")

    tiers = [
        ("High", "125", "Mandatory review", "58% - 94%"),
        ("Medium", "125", "Overflow capacity", "40% - 58%"),
        ("Low", "250", "Standard submission", "11% - 40%"),
    ]
    for tier, count, priority, prob in tiers:
        row = table.add_row().cells
        row[0].text = tier
        row[1].text = count
        row[2].text = priority
        row[3].text = prob
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    doc.add_paragraph()

    # Section 3: What Drives Denials
    add_heading(doc, "3. Key Findings -- What Drives Denials", level=1)
    add_para(doc, "The model identified clear, actionable patterns that increase denial risk:", space_after=4)
    add_bullet(doc, "Medicaid MCO claims: denial rate of 31%, nearly double the commercial rate of 15%.")
    add_bullet(doc, "Inpatient visits: denial rate of 31%, compared to 18-20% for other visit types.")
    add_bullet(doc, "Missing prior authorization alone raises denial risk to 47%.")
    add_bullet(doc, "Missing documentation raises denial risk to 39%.")
    add_bullet(doc, "When both authorization and documentation are missing together, denial risk jumps to 74%.")
    add_para(doc,
        "These are not mysterious statistical signals -- they are concrete operational gaps that the "
        "review team can fix before submission. The model simply surfaces them automatically and ranks "
        "claims by combined risk.",
        space_after=8)

    # Section 4: Model Results
    add_heading(doc, "4. Model Selection Results", level=1)
    add_para(doc,
        "We tested five model architectures in an active learning experiment to find the best performer. "
        "The selected model is Calibrated Logistic Regression, which delivers the same denial capture rate "
        "as other top performers while providing the most reliable probability estimates for tier assignment.",
        space_after=6)

    # Experiment table
    add_para(doc, "Experiment Leaderboard:", bold=True, space_after=4)
    table2 = doc.add_table(rows=1, cols=5)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Model"
    hdr2[1].text = "Denial Capture"
    hdr2[2].text = "ROC-AUC"
    hdr2[3].text = "Brier Score"
    hdr2[4].text = "Verdict"
    for cell in hdr2:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")

    experiments = [
        ("Calibrated Logistic Regression", "49.5%", "0.710", "0.137", "SELECTED"),
        ("Logistic Regression (baseline)", "49.5%", "0.711", "0.209", "Runner-up"),
        ("Logistic Regression + interactions", "49.5%", "0.707", "0.210", "No gain"),
        ("Random Forest", "39.8%", "0.653", "0.155", "Underperforms"),
        ("XGBoost", "35.9%", "0.615", "0.180", "Overfits"),
    ]
    for model, capture, auc, brier, verdict in experiments:
        row = table2.add_row().cells
        row[0].text = model
        row[1].text = capture
        row[2].text = auc
        row[3].text = brier
        row[4].text = verdict
        if verdict == "SELECTED":
            for paragraph in row[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            for paragraph in row[4].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    table2.alignment = WD_TABLE_ALIGNMENT.LEFT
    doc.add_paragraph()

    add_para(doc,
        "Why Logistic Regression won: the synthetic data was generated by simple additive rules -- each risk "
        "factor adds a fixed amount of risk. Tree-based models (Random Forest, XGBoost) overfit to spurious "
        "interactions at this sample size, while Logistic Regression correctly models the additive pattern. "
        "Calibration additionally improves probability reliability by 34%, making the High/Medium/Low tiers "
        "more trustworthy for operations staff.",
        space_after=8)

    # Section 5: GenAI Explanations
    add_heading(doc, "5. GenAI Explanation Engine", level=1)
    add_para(doc,
        "For the top 10 highest-risk claims, a GenAI engine (Ollama Cloud, Gemma 4) produces "
        "plain-English explanations that tell the analyst exactly why the claim is flagged and what "
        "corrective action to take. Every output is validated for accuracy, includes a mandatory "
        "uncertainty disclaimer, and falls back to a deterministic template if the API is unavailable.",
        space_after=6)
    add_bullet(doc, "Example for a high-risk claim: 'Missing required prior authorization. Confirm and obtain the required prior authorization before submission.'")
    add_bullet(doc, "Example for a low-risk claim: 'No actionable pre-submission risk flags. Routine submission with standard verification is recommended.'")
    add_para(doc,
        "All LLM calls are audited for token usage, latency, and response quality. The latest production "
        "run processed 10 claims with 100% validation pass rate, 4,479 total tokens, and zero hallucination "
        "warnings.",
        space_after=8)

    # Section 6: Safety & Compliance
    add_heading(doc, "6. Safety, Auditability, and Compliance", level=1)
    add_numbered(doc, "Leakage prevention: the target variable (is_denied) and post-outcome data are programmatically excluded from model inputs.")
    add_numbered(doc, "Deterministic fallback: if the GenAI API is unavailable, the system generates clean, pre-approved explanations without any external dependency.")
    add_numbered(doc, "HIPAA-safe audit logs: all LLM observability data (tokens, latency, quality flags) is stored in structured local JSON files -- no telemetry leaves the machine.")
    add_numbered(doc, "Reproducible experiments: every model version is fully versioned with parameters, metrics, and serialized artifacts, enabling rollback if needed.")
    doc.add_paragraph()

    # Section 7: Limitations & Next Steps
    add_heading(doc, "7. Limitations and Recommendations", level=1)
    add_para(doc, "Limitations:", bold=True, space_after=4)
    add_bullet(doc, "The data is synthetic -- real claims include ICD/CPT codes, 835 remittance data, and clinical narratives that would significantly improve the model.")
    add_bullet(doc, "Denial rates increased from ~20% in training to 26% in the test period, suggesting temporal drift that requires ongoing monitoring.")
    add_bullet(doc, "The model cannot capture complex payer-specific rules without payer-specific training data.")
    add_para(doc, "Recommendations for production:", bold=True, space_after=4)
    add_bullet(doc, "Deploy the calibrated Logistic Regression as the production model immediately.")
    add_bullet(doc, "Retrain on real claims data with ICD/CPT codes and line-item detail within the first 90 days.")
    add_bullet(doc, "Implement drift monitoring (denial rate shifts, feature distribution changes) with monthly reviews.")
    add_bullet(doc, "Add dollar-weighted optimization so the queue prioritizes high-dollar denials, not just high-count denials.")
    add_bullet(doc, "Establish an A/B test comparing model-assisted vs. manual triage over 3-6 months to measure actual dollar recovery.")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runf = footer.add_run("--- End of Document ---")
    runf.italic = True
    runf.font.size = Pt(10)
    runf.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Save
    out_path = os.path.join(os.path.dirname(__file__), "Project_documentation.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == "__main__":
    build_document()
